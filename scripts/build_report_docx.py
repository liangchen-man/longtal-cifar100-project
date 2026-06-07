from __future__ import annotations

from pathlib import Path
import csv

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
FIG_DIR = REPORT_DIR / "figures" / "generated"
RESULTS_DIR = ROOT / "results"
OUT_PATH = REPORT_DIR / "机器学习课程作业_图像长尾分类实验报告.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths_dxa[idx])


def set_cell_margins(table, margin: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for side in ["top", "left", "bottom", "right"]:
        node = tbl_cell_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), "80" if side in {"top", "bottom"} else str(margin))
        node.set(qn("w:type"), "dxa")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    if level == 1:
        p.style = doc.styles["Heading 1"]
    elif level == 2:
        p.style = doc.styles["Heading 2"]
    else:
        p.style = doc.styles["Heading 3"]
    p.add_run(text)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_cell_margins(table)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], GRAY_FILL)
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_picture_if_exists(doc: Document, path: Path, caption: str, width: float = 6.1) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = doc.styles["Caption"]
    cap.add_run(caption)


def load_summary_rows() -> list[list[str]]:
    path = RESULTS_DIR / "summary.csv"
    if not path.exists():
        return []
    method_order = [
        "HOG+PCA+LinearSVM",
        "ce_resnet32",
        "ldam_drw_resnet32",
        "ldam_drw_resnet32_s1",
        "balanced_softmax_resnet32",
        "logit_adjusted_resnet32",
        "class_balanced_focal_resnet32",
    ]
    method_labels = {
        "ce_resnet32": "ResNet-32+CE",
        "ldam_drw_resnet32": "LDAM-DRW s=30",
        "ldam_drw_resnet32_s1": "LDAM-DRW s=1",
        "balanced_softmax_resnet32": "Balanced Softmax",
        "logit_adjusted_resnet32": "Logit Adjustment τ=1",
        "class_balanced_focal_resnet32": "CB-Focal",
    }
    split_order = {"train": 0, "balanced_test": 1, "imbalanced_test": 2}
    rows: list[list[str]] = []
    with path.open("r", encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        raw_rows = [
            row for row in reader
            if row.get("dataset") == "cifar100" and row.get("method") in method_order
        ]
        raw_rows.sort(key=lambda row: (method_order.index(row["method"]), split_order.get(row.get("split", ""), 99)))
        for row in raw_rows:
            if row.get("dataset") != "cifar100":
                continue
            rows.append(
                [
                    method_labels.get(row.get("method", ""), row.get("method", "")),
                    row.get("split", ""),
                    f"{float(row.get('top1') or 0):.2f}",
                    f"{float(row.get('top5') or 0):.2f}",
                    f"{float(row.get('macro_f1') or 0):.2f}",
                    f"{float(row.get('balanced_acc') or 0):.2f}",
                    f"{float(row.get('few_acc') or 0):.2f}",
                ]
            )
    return rows


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.font.size = Pt(9)
    caption.font.italic = True


def build_docx() -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CIFAR-100-LT 图像长尾分类实验报告")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("班级：__________    姓名：__________    学号：__________    日期：2026 年 6 月")

    add_heading(doc, "摘要", 1)
    add_body(
        doc,
        "本项目围绕 CIFAR-100-LT 长尾图像分类任务，构建了 exp/step 数据协议、传统机器学习基线、"
        "ResNet-32 神经网络基线以及 LDAM-DRW、Balanced Softmax、Logit Adjustment、Class-Balanced Focal "
        "等长尾改进策略。评估覆盖训练集、平衡测试集和同分布不平衡测试集，并报告 Top-1、Top-5、"
        "Macro-F1、Balanced Accuracy、Many/Medium/Few Accuracy、ECE、逐类准确率与混淆矩阵。"
    )
    add_body(
        doc,
        "本次已完成 CIFAR-100-LT exp IR=100 正式训练与评估。CE 在不平衡测试集 Top-1 "
        "达到 64.71%，但平衡测试 Few Acc 仅 9.63%；Balanced Softmax / Logit Adjustment τ=1 将平衡测试 "
        "Top-1 提升到 44.89%，Macro-F1 提升到 43.93%，Few Acc 提升到 27.27%，是本轮补跑中综合最好的改进策略。"
    )

    add_heading(doc, "1. 任务与数据集构建", 1)
    add_body(
        doc,
        "CIFAR-100-LT 从 CIFAR-100 训练集按类别抽样构造长尾分布。exp 协议令第 i 类样本数按指数衰减，"
        "step 协议令前半类别为头部、后半类别为尾部。IR=100 时，训练集头部类别约 500 张，尾部类别约 5 张。"
    )
    add_bullets(
        doc,
        [
            "训练集：按 exp 或 step 方式抽样得到长尾训练集。",
            "平衡测试集：使用原始 CIFAR-100 测试集，每类 100 张。",
            "不平衡测试集：按训练集同分布从测试集抽样，用于模拟真实部署分布。",
        ],
    )
    add_picture_if_exists(doc, FIG_DIR / "distribution_exp_ir100.png", "图 1  exp IR=100 训练集与不平衡测试集类别分布")
    add_picture_if_exists(doc, FIG_DIR / "distribution_step_ir100.png", "图 2  step IR=100 训练集与不平衡测试集类别分布")

    add_heading(doc, "2. 模型与长尾改进策略", 1)
    add_heading(doc, "2.1 传统机器学习：HOG + PCA + Linear SVM", 2)
    add_body(
        doc,
        "传统基线对图像提取 HOG 特征，经过 StandardScaler 和 PCA 降维后，使用带 class_weight=balanced 的线性 SVM 分类。"
        "该方法可解释、训练成本低，但对 CIFAR-100 的高层语义类别表达有限。"
    )
    add_heading(doc, "2.2 神经网络：ResNet-32 + CE", 2)
    add_body(
        doc,
        "神经网络基线采用 CIFAR 风格 ResNet-32，使用随机裁剪、水平翻转和标准化增强。标准 CE 容易被头部类别主导，"
        "因此需要配合长尾学习策略观察尾部类别改善。"
    )
    add_heading(doc, "2.3 改进策略", 2)
    add_bullets(
        doc,
        [
            "LDAM-DRW：尾部类别使用更大 margin，训练后期再启用 effective-number 权重。",
            "Balanced Softmax：将训练集类别先验加入 softmax，修正标签分布偏移。",
            "Logit Adjustment：在 logits 上加入类别先验校正，提高平衡测试与校准表现。",
            "Class-Balanced Focal：结合类别权重与难例调制，降低头部易样本的梯度主导。",
        ],
    )

    add_heading(doc, "3. 实验矩阵与评价指标", 1)
    add_table(
        doc,
        ["方法", "类型", "配置文件", "目的"],
        [
            ["HOG+PCA+LinearSVM", "传统机器学习", "configs/exp_ir100.yaml", "课程传统方法要求"],
            ["ResNet-32+CE", "神经网络基线", "configs/ce_resnet32_ir100.yaml", "观察长尾偏置"],
            ["ResNet-32+LDAM-DRW", "主改进方法", "configs/ldam_drw_resnet32_ir100.yaml", "原始 s=30 设置"],
            ["ResNet-32+LDAM-DRW(s=1)", "调参改进", "configs/ldam_drw_resnet32_ir100_s1.yaml", "提升尾部类和宏平均指标"],
            ["ResNet-32+Balanced Softmax", "分布修正", "configs/balanced_softmax_resnet32_ir100.yaml", "修正类别先验偏移"],
            ["ResNet-32+Logit Adjustment", "后验校正", "configs/logit_adjusted_resnet32_ir100.yaml", "改善平衡测试和校准"],
            ["ResNet-32+CB-Focal", "权重与难例调制", "configs/class_balanced_focal_resnet32_ir100.yaml", "降低头部易例主导"],
        ],
        [2100, 1700, 3300, 2260],
    )
    add_body(
        doc,
        "评价指标包括 Top-1、Top-5、Macro-F1、Balanced Accuracy、Many/Medium/Few Acc、ECE、逐类准确率、混淆矩阵和 reliability diagram。"
        "其中 Macro-F1、Balanced Acc 和 Few Acc 对长尾识别更敏感，能避免 overall accuracy 掩盖尾部失败。"
    )

    add_heading(doc, "4. CIFAR-100-LT 正式实验结果", 1)
    add_body(
        doc,
        "本次使用项目内 .conda-cuda 环境运行 200 epoch 神经网络实验，并使用本地 CIFAR-100 数据构造 exp IR=100 "
        "长尾训练集。表中同时保留原始 s=30 LDAM-DRW 未收敛结果、s=1 调参结果，以及 Balanced Softmax、Logit Adjustment 和 CB-Focal 正式补跑结果。"
    )
    summary_rows = load_summary_rows()
    if summary_rows:
        add_table(
            doc,
            ["方法", "Split", "Top-1", "Top-5", "Macro-F1", "Balanced Acc", "Few Acc"],
            summary_rows,
            [1900, 1600, 950, 950, 1150, 1400, 1250],
        )
    add_picture_if_exists(doc, FIG_DIR / "comparison_top1.png", "图 3  不同测试协议下 Top-1 对比")
    add_picture_if_exists(doc, FIG_DIR / "comparison_macro_f1.png", "图 4  Macro-F1 对比")
    add_picture_if_exists(doc, FIG_DIR / "comparison_few_acc.png", "图 5  Few Acc 对比")
    add_picture_if_exists(doc, FIG_DIR / "per_class_ce_resnet32_ir100_balanced_test.png", "图 6  CE 平衡测试集逐类准确率")
    add_picture_if_exists(doc, FIG_DIR / "per_class_ldam_drw_resnet32_ir100_s1_balanced_test.png", "图 7  LDAM-DRW(s=1) 平衡测试集逐类准确率")
    add_picture_if_exists(doc, FIG_DIR / "ldam_scale_ablation.png", "图 8  LDAM scale 30 epoch 消融")
    add_picture_if_exists(doc, FIG_DIR / "confusion_balanced_softmax_resnet32_ir100_balanced_test.png", "图 9  Balanced Softmax 平衡测试集混淆矩阵")
    add_picture_if_exists(doc, FIG_DIR / "reliability_balanced_softmax_resnet32_ir100_balanced_test.png", "图 10  Balanced Softmax 平衡测试集 reliability diagram")

    add_heading(doc, "4.1 LDAM scale 诊断", 2)
    add_body(
        doc,
        "30 epoch 短跑显示，LDAM 的 scale 对优化稳定性非常敏感：s=0.5/1/2 可以正常学习，s>=5 迅速退化到接近随机水平。"
    )
    add_table(
        doc,
        ["s", "Balanced Top-1", "Macro-F1", "Few Acc", "ECE"],
        [
            ["0.5", "34.50", "29.31", "9.20", "36.63"],
            ["1", "35.90", "31.34", "10.73", "12.97"],
            ["2", "34.94", "30.69", "12.40", "12.12"],
            ["5", "4.44", "1.45", "0.03", "2.91"],
            ["10", "1.00", "0.02", "0.00", "0.42"],
            ["30", "1.01", "0.04", "0.00", "21.23"],
        ],
        [1200, 1900, 1700, 1500, 1400],
    )

    add_heading(doc, "5. 分析与结论", 1)
    add_numbered(
        doc,
        [
            "不平衡测试集 overall Top-1 可能高于平衡测试集，因为测试分布更偏向高频类别；这不代表模型真正解决了尾部识别。",
            "CE 的不平衡测试 Top-1 为 64.71%，但平衡测试 Few Acc 只有 9.63%，说明尾部类别仍然薄弱。",
            "原始 s=30 LDAM-DRW 在本设置中未收敛，平衡测试 Top-1 只有 1.00%，不能作为有效改进结论。",
            "调参版 LDAM-DRW(s=1) 将平衡测试 Top-1 提升到 44.60%，Few Acc 提升到 19.87%，更符合长尾分类目标。",
            "Balanced Softmax / Logit Adjustment τ=1 将平衡测试 Top-1 提升到 44.89%，Few Acc 提升到 27.27%，是本轮补跑中综合最好的策略；两者结果一致来自当前实现中相同的类别先验校正形式。",
            "Class-Balanced Focal 的 Few Acc 为 19.13%，但平衡测试 Top-1 只有 21.55%，说明过强重加权会明显损害整体判别能力。",
        ],
    )
    add_body(
        doc,
        "本项目已完成课程要求并扩展为完整可复现实验框架。结果说明，长尾分类不能只看 overall accuracy；"
        "平衡测试集、Macro-F1、Balanced Acc、Few Acc 与逐类准确率更能揭示尾部类别是否真正被学到。"
    )

    add_heading(doc, "6. 复现命令", 1)
    add_body(doc, "当前环境快速验证：")
    add_body(doc, r".\scripts\run_smoke_pipeline.ps1")
    add_body(doc, "完整 CIFAR-100-LT 训练：")
    add_body(doc, r".\scripts\run_full_experiments.ps1")

    add_heading(doc, "7. 参考文献", 1)
    references = [
        "Cao K., Wei C., Gaidon A., Arechiga N., Ma T. Learning Imbalanced Datasets with Label-Distribution-Aware Margin Loss. NeurIPS, 2019. https://arxiv.org/abs/1906.07413",
        "Cui Y., Jia M., Lin T.-Y., Song Y., Belongie S. Class-Balanced Loss Based on Effective Number of Samples. CVPR, 2019. https://arxiv.org/abs/1901.05555",
        "Lin T.-Y., Goyal P., Girshick R., He K., Dollar P. Focal Loss for Dense Object Detection. ICCV, 2017. https://arxiv.org/abs/1708.02002",
        "Menon A. K., Jayasumana S., Rawat A. S., Jain H., Veit A., Kumar S. Long-tail Learning via Logit Adjustment. ICLR, 2021. https://openreview.net/forum?id=37nvvqkCo5",
        "Ren J., Yu C., Sheng S., Ma X., Zhao H., Yi S., Li H. Balanced Meta-Softmax for Long-Tailed Visual Recognition. NeurIPS, 2020. https://arxiv.org/abs/2007.10740",
        "Kang B., Xie S., Rohrbach M., Yan Z., Gordo A., Feng J., Kalantidis Y. Decoupling Representation and Classifier for Long-Tailed Recognition. ICLR, 2020. https://openreview.net/forum?id=r1gRTCVFvB",
        "Wang X., Lian L., Miao Z., Liu Z., Yu S. X. Long-tailed Recognition by Routing Diverse Distribution-Aware Experts. ICLR, 2021. https://openreview.net/forum?id=D9I3drBz4UC",
        "Cui J., Liu S., Tian Z., Zhong Z., Jia J. ResLT: Residual Learning for Long-Tailed Recognition. IEEE TPAMI, 2023. https://arxiv.org/abs/2101.10633",
    ]
    for idx, reference in enumerate(references, start=1):
        add_body(doc, f"[{idx}] {reference}")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_docx()
